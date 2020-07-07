# -*- coding: utf-8 -*-
"""
Created on Mon Jun 22 20:24:52 2020

@author: Precious
"""

from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from PIL import Image
import docx
from docx.shared import Inches
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import seaborn as sns
import pdb


def wordcreator(issue,medication):
    try:
    
        wc = WordCloud(stopwords=stop_words,
               background_color="white",
               colormap="Dark2",
               max_font_size=300,
               max_words = 30,
               random_state=42)
    
    
        condition =issue
        drugname =medication
        #original dataset that has not been summarised
        realdata = data[['drugName', 'condition','rating','usefulCount']]
        realdata = realdata[data.condition == condition]
        ratingplot = sns.boxplot(x='drugName',y='rating',hue='drugName',data=realdata).set_title('Boxplot of Medications')
        
        #summarised data that has been cleaned 
        alldrugs = data_combined[(data_combined.Condition == condition)]
        requesteddrug = data_combined[(data_combined.Condition ==condition) & (data_combined.drugName == drugname)]
        relateddrugs =  data_combined[(data_combined.Condition ==condition) & (data_combined.drugName != drugname)]
        max_average_rating = list(alldrugs.drugName[alldrugs.averageRating == alldrugs.averageRating.max()])[0]
        max_usefulTo =list(alldrugs.drugName[alldrugs.usefulTo == alldrugs.usefulTo.max()])[0]
        col1 = (alldrugs["usefulTo"]/100).astype(str)
        col2 = alldrugs["averageRating"].astype(str)
        #scatplot = sns.regplot(x='usefulTo',y='averageRating',data=alldrugs,fit_reg = False)
        
        alldrugs['Rank'] = (col2+col1).rank(method='dense', ascending=False).astype(int)
        sorted_rank =alldrugs.sort_values('Rank')
        sorted_rank = sorted_rank[['drugName','averageRating','usefulTo','Rank']]
        alldrugs.drop(['Rank'],axis=1)
        
    
        filename = 'condition_medication_report.pdf'
        title = "Report on {0} for {1}".format(drugname,condition)
        subtitle =['Word Cloud on {0}'.format(drugname),
                   'Statistics Summary on {0}'.format(drugname),
                   'Related Medication Others Used for {0}'.format(condition),
                   'Word Cloud for other related Medications for {0}'.format(condition),
                   'Summary Statistics of other related medications',
                   'Best Medication for {0} according to reviews'.format(condition)]
        message = ["Average Rating of  : {0}".format(list(requesteddrug.averageRating)), 
                   "Found Useful To(persons) : {0}".format(list(requesteddrug.usefulTo)),
                   "List of other drugs: {0}".format(list(relateddrugs.drugName.unique())),
                   #"Best Medication for {0} is {1}".format(condition,sorted_rank[0]),
                   "Best medication  based on Average rating : {0}".format(max_average_rating),
                   "best Medication by number of person(s) that found it useful: {0}".format(max_usefulTo)]
        
        
        doc = docx.Document()
        doc.add_heading(title,level=1)
        doc.add_heading(subtitle[0],level=3)
        wcdata = data_combined.chunk[(data_combined.Condition == condition) & (data_combined.drugName==drugname)]
        wc_topword= top_word[(condition,drugname)]
        wc.generate(str(wc_topword))
        wc.to_file('wc.jpeg')
        figure=Image.open('wc.jpeg')
        figure=figure.resize((200,100))
        figure.save('wc_resized.jpeg')
        doc.add_picture('wc_resized.jpeg',width=Inches(2.08),height=Inches(1.04))
        doc.add_heading(subtitle[1],level=3)
        doc.add_paragraph(message[0])
        doc.add_paragraph(message[1])
        if (relateddrugs.empty !=True):
            doc.add_heading(subtitle[2],level=3)
            doc.add_paragraph(message[2])
            doc.add_heading(subtitle[3],level=3)
            doc.add_paragraph()
            unique_related_drug = list(relateddrugs.drugName.unique())
        
            for drug in unique_related_drug:
                t= top_word[(condition,drug)]
                wc.generate(str(t))
                wc.to_file('fig.jpeg')
                doc.add_picture('fig.jpeg',width = Inches(2.08),height = Inches(1.04))
                doc.add_paragraph("figure: "+drug)
            
            doc.add_heading(subtitle[4],level =3)
            ratingfig = ratingplot.get_figure()
            ratingfig.savefig('plot.jpeg')
            doc.add_picture('plot.jpeg',width=Inches(5),height=Inches(2.5))
            doc.add_paragraph()
            '''scatplotfig = scatplot.get_figure()
            scatplotfig.savefig('plot2.jpeg')
            doc.add_picture('plot2.jpeg',width=Inches(5),height=Inches(2.5))
            '''
            
            doc.add_heading(subtitle[5],level =3)
            doc.add_paragraph(message[3])
            doc.add_paragraph(message[4])
            doc.add_paragraph()
            doc.add_heading("Ranked Medications for {0}".format(condition),level =3)
            table2 = doc.add_table(rows=sorted_rank.shape[0], cols=4)
            
            for i  in range(sorted_rank.shape[0]):
                row_cells = table2.add_row().cells
                row_cells[0].text = sorted_rank.iloc[i,0]
                row_cells[1].text = str(sorted_rank.iloc[i,1])
                row_cells[2].text = str(sorted_rank.iloc[i,2])
                row_cells[3].text = str(sorted_rank.iloc[i,3])
                hdr_cells = table2.rows[0].cells
                hdr_cells[0].text = 'drugName'
                hdr_cells[1].text = 'averageRating'
                hdr_cells[2].text = 'usefulTo'
                hdr_cells[3].text = 'Rank'
                
        doc.save("{0}_report.docx".format(drugname))
                
    except:
        print("Something went wrong, {0} or/and {1} were not found in the dataset".format(condition,drugname))
        
    